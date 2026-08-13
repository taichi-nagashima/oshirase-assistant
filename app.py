import io
import os
from datetime import datetime

import streamlit as st
from docxtpl import DocxTemplate
from docx import Document
from pydantic import BaseModel, Field
from google import genai
from google.genai import types


# =============================================================================
# Page configuration
# =============================================================================
st.set_page_config(
    layout="wide",
    page_title="お便りメーカー",
    page_icon="✦",
    initial_sidebar_state="collapsed",
)


# =============================================================================
# Constants
# =============================================================================
TEMPLATE_DIR = "templates"
os.makedirs(TEMPLATE_DIR, exist_ok=True)

AUDIENCE_OPTIONS = {
    "student": "生徒向け",
    "parent": "保護者向け",
    "grade": "学年全体向け",
}
DOCUMENT_TYPE_OPTIONS = {
    "class_newsletter": "学級通信",
    "event_notice": "行事のお知らせ",
    "club_notice": "部活動の連絡",
}
LENGTH_OPTIONS = {
    "short": "約400字",
    "long": "約1000字",
}
GEMINI_MODEL = "gemini-2.5-flash"

DAY_KEYS = ["mon", "tue", "wed", "thu", "fri"]
DAY_LABELS = {
    "mon": "月",
    "tue": "火",
    "wed": "水",
    "thu": "木",
    "fri": "金",
}


# =============================================================================
# Default template
# =============================================================================
def ensure_default_template() -> None:
    default_path = os.path.join(TEMPLATE_DIR, "標準テンプレート.docx")
    if not os.path.exists(default_path):
        doc = Document()
        doc.add_heading("{{ title }}", level=0)
        doc.add_paragraph("{{ greeting }}")
        doc.add_paragraph("{{ body }}")
        doc.add_heading("開催概要", level=1)
        doc.add_paragraph("日時: {{ event_date }} {{ event_time }}")
        doc.add_paragraph("場所: {{ event_place }}")
        doc.add_paragraph("持ち物: {{ event_items }}")
        doc.add_paragraph("{{ closing }}")
        doc.save(default_path)


ensure_default_template()


# =============================================================================
# Data schema
# =============================================================================
class NewsletterData(BaseModel):
    title: str = Field(description="お便りのタイトル")
    subtitle: str = Field(description="ヘッダー用サブタイトル")
    date_str: str = Field(description="発行日")
    greeting: str = Field(description="導入の温かい挨拶文")
    body: str = Field(description="本文")
    parent_note: str = Field(description="保護者への注意書き")
    closing: str = Field(description="結びの言葉")

    mon_1: str = ""
    mon_2: str = ""
    mon_3: str = ""
    mon_4: str = ""
    mon_5: str = ""
    mon_6: str = ""

    tue_1: str = ""
    tue_2: str = ""
    tue_3: str = ""
    tue_4: str = ""
    tue_5: str = ""
    tue_6: str = ""

    wed_1: str = ""
    wed_2: str = ""
    wed_3: str = ""
    wed_4: str = ""
    wed_5: str = ""
    wed_6: str = ""

    thu_1: str = ""
    thu_2: str = ""
    thu_3: str = ""
    thu_4: str = ""
    thu_5: str = ""
    thu_6: str = ""

    fri_1: str = ""
    fri_2: str = ""
    fri_3: str = ""
    fri_4: str = ""
    fri_5: str = ""
    fri_6: str = ""

    urls: list[str] = Field(default_factory=list)


# =============================================================================
# UI
# =============================================================================
def apply_apple_styles() -> None:
    st.markdown(
        """
        <style>
        :root {
            --ink: #1d1d1f;
            --muted: #6e6e73;
            --line: #e5e5e7;
            --soft: #f5f5f7;
            --card: #ffffff;
            --accent: #0071e3;
        }

        html, body, [class*="css"] {
            font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display",
                         "SF Pro Text", "Helvetica Neue", Arial, sans-serif;
        }

        .stApp {
            background:
                radial-gradient(circle at 50% -10%, rgba(0,113,227,.07), transparent 32rem),
                #f5f5f7;
            color: var(--ink);
        }

        .block-container {
            max-width: 1180px;
            padding-top: 2.5rem;
            padding-bottom: 5rem;
        }

        header[data-testid="stHeader"] {
            background: rgba(245,245,247,.72);
            backdrop-filter: blur(18px);
        }

        /* Hide Streamlit's default decoration without hiding controls. */
        #MainMenu { visibility: hidden; }
        footer { visibility: hidden; }

        .hero {
            text-align: center;
            padding: 1.4rem 0 2.3rem;
        }

        .eyebrow {
            display: inline-flex;
            align-items: center;
            gap: .4rem;
            padding: .42rem .75rem;
            border: 1px solid rgba(0,113,227,.13);
            border-radius: 999px;
            background: rgba(255,255,255,.68);
            color: var(--accent);
            font-size: .76rem;
            font-weight: 650;
            letter-spacing: .03em;
        }

        .hero h1 {
            margin: .9rem 0 .35rem;
            font-size: clamp(2.25rem, 5vw, 4.1rem);
            line-height: 1.02;
            letter-spacing: -.055em;
            font-weight: 750;
        }

        .hero p {
            margin: 0 auto;
            max-width: 650px;
            color: var(--muted);
            font-size: 1.02rem;
            line-height: 1.7;
        }

        .section-title {
            margin: 2.1rem 0 .75rem;
            font-size: 1.15rem;
            font-weight: 700;
            letter-spacing: -.02em;
        }

        .section-caption {
            margin: -.35rem 0 1rem;
            color: var(--muted);
            font-size: .88rem;
        }

        .card {
            background: rgba(255,255,255,.88);
            border: 1px solid rgba(0,0,0,.055);
            border-radius: 24px;
            padding: 1.35rem;
            box-shadow: 0 10px 35px rgba(0,0,0,.045);
        }

        .mini-label {
            color: var(--muted);
            font-size: .76rem;
            font-weight: 650;
            letter-spacing: .04em;
            margin-bottom: .35rem;
        }

        .status-pill {
            display: inline-flex;
            align-items: center;
            gap: .45rem;
            border-radius: 999px;
            padding: .35rem .65rem;
            font-size: .75rem;
            font-weight: 600;
            background: #f0f0f2;
            color: #6e6e73;
        }

        .status-dot {
            width: 7px;
            height: 7px;
            border-radius: 50%;
            background: #34c759;
        }

        .preview-paper {
            background: #fff;
            border: 1px solid var(--line);
            border-radius: 18px;
            padding: 2rem;
            box-shadow: 0 12px 30px rgba(0,0,0,.05);
        }

        .preview-title {
            font-size: 1.75rem;
            line-height: 1.25;
            letter-spacing: -.035em;
            font-weight: 750;
            margin-bottom: .25rem;
        }

        .preview-subtitle {
            color: var(--muted);
            font-size: .88rem;
            margin-bottom: 1.35rem;
        }

        .preview-heading {
            font-weight: 700;
            margin: 1.2rem 0 .45rem;
        }

        .preview-body {
            white-space: pre-wrap;
            line-height: 1.9;
            color: #303034;
        }

        /* Inputs */
        div[data-baseweb="select"] > div,
        .stTextInput > div > div,
        .stTextArea > div > textarea {
            border-radius: 14px;
            border-color: #d9d9dc;
            background: rgba(255,255,255,.9);
        }

        .stTextArea textarea {
            min-height: 180px;
            line-height: 1.7;
        }

        .stButton > button,
        .stDownloadButton > button {
            border-radius: 999px;
            min-height: 2.75rem;
            font-weight: 650;
            border: 1px solid rgba(0,0,0,.08);
            transition: transform .15s ease, box-shadow .15s ease;
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            transform: translateY(-1px);
            box-shadow: 0 8px 22px rgba(0,0,0,.09);
        }

        .generate-button button {
            width: 100%;
            min-height: 3.35rem;
            font-size: 1rem;
        }

        div[data-testid="stExpander"] {
            border: 1px solid rgba(0,0,0,.07);
            border-radius: 18px;
            background: rgba(255,255,255,.65);
        }

        .timetable-note {
            color: var(--muted);
            font-size: .82rem;
            margin-bottom: .9rem;
        }

        .divider {
            height: 1px;
            background: var(--line);
            margin: 2.2rem 0;
        }

        @media (max-width: 768px) {
            .block-container {
                padding-left: 1rem;
                padding-right: 1rem;
                padding-top: 1.2rem;
            }
            .hero {
                padding-bottom: 1.3rem;
            }
            .card, .preview-paper {
                border-radius: 18px;
                padding: 1rem;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# =============================================================================
# Logic
# =============================================================================
def get_api_key() -> str | None:
    for key_name in ("GEMINI_API_KEY", "GOOGLE_API_KEY"):
        if value := os.environ.get(key_name):
            return value.strip()

    try:
        for key_name in ("GEMINI_API_KEY", "GOOGLE_API_KEY"):
            if key_name in st.secrets:
                return str(st.secrets[key_name]).strip()
    except Exception:
        pass

    return None


def get_available_templates() -> list[str]:
    return sorted(
        [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx")]
    )


def generate_newsletter_data(
    api_key: str, prompt: str, use_web_search: bool
) -> NewsletterData:
    client = genai.Client(api_key=api_key)
    tools = [types.Tool(google_search=types.GoogleSearch())] if use_web_search else None

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            tools=tools,
            response_mime_type="application/json",
            response_schema=NewsletterData,
            temperature=0.7,
        ),
    )

    urls = []
    if use_web_search and response.candidates:
        metadata = getattr(response.candidates[0], "grounding_metadata", None)
        if metadata and getattr(metadata, "grounding_chunks", None):
            for chunk in metadata.grounding_chunks:
                web = getattr(chunk, "web", None)
                if web and web.uri and web.uri not in urls:
                    urls.append(web.uri)

    result_data = NewsletterData.model_validate_json(response.text)
    if urls:
        result_data.urls = urls

    return result_data


def render_docx(template_name: str, data: NewsletterData) -> bytes:
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    doc = DocxTemplate(template_path)
    doc.render(data.model_dump())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def timetable_from_state() -> dict[str, list[str]]:
    return st.session_state.get(
        "timetable_input",
        {key: [""] * 6 for key in DAY_KEYS},
    )


def timetable_text(timetable: dict[str, list[str]]) -> str:
    lines = []
    for key in DAY_KEYS:
        values = timetable.get(key, [""] * 6)
        subjects = [v.strip() if v.strip() else "（なし）" for v in values]
        lines.append(f"{DAY_LABELS[key]}曜日: " + " / ".join(subjects))
    return "\n".join(lines)


def main() -> None:
    apply_apple_styles()

    if "doc_data" not in st.session_state:
        st.session_state.doc_data = None

    if "timetable_input" not in st.session_state:
        st.session_state.timetable_input = {
            key: [""] * 6 for key in DAY_KEYS
        }

    api_key = get_api_key()
    templates = get_available_templates()

    # -------------------------------------------------------------------------
    # Hero
    # -------------------------------------------------------------------------
    st.markdown(
        """
        <div class="hero">
            <div class="eyebrow">✦ Teacher's writing assistant</div>
            <h1>お便りを、<br>もっと簡単に。</h1>
            <p>
                伝えたいことを入力するだけ。<br>
                AIが、先生らしい温かいお便りに整えます。
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -------------------------------------------------------------------------
    # Settings card
    # -------------------------------------------------------------------------
    st.markdown('<div class="section-title">1. 文書の設定</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-caption">誰に、どんなお便りを届けるかを選びます。</div>',
        unsafe_allow_html=True,
    )

    with st.container(border=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            audience = st.selectbox(
                "対象読者",
                list(AUDIENCE_OPTIONS.keys()),
                format_func=lambda k: AUDIENCE_OPTIONS[k],
            )
        with c2:
            doc_type = st.selectbox(
                "文書の種類",
                list(DOCUMENT_TYPE_OPTIONS.keys()),
                format_func=lambda k: DOCUMENT_TYPE_OPTIONS[k],
            )
        with c3:
            length = st.selectbox(
                "目安の分量",
                list(LENGTH_OPTIONS.keys()),
                format_func=lambda k: LENGTH_OPTIONS[k],
            )

    # -------------------------------------------------------------------------
    # Main input
    # -------------------------------------------------------------------------
    st.markdown('<div class="section-title">2. 伝えたいこと</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-caption">箇条書きでも、思いついたままでも大丈夫です。</div>',
        unsafe_allow_html=True,
    )

    draft_notes = st.text_area(
        "伝えたい内容",
        height=190,
        label_visibility="collapsed",
        placeholder=(
            "例：\n"
            "・今週は係活動を頑張っていた\n"
            "・子どもたちから「もっとクラスを良くしたい」という声が出た\n"
            "・来週は校外学習があります\n"
            "・保護者に持ち物を伝えたい"
        ),
    )

    # -------------------------------------------------------------------------
    # Timetable
    # -------------------------------------------------------------------------
    st.markdown('<div class="section-title">3. 来週の時間割</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-caption">必要な場合だけ入力してください。空欄でも生成できます。</div>',
        unsafe_allow_html=True,
    )

    with st.expander("時間割を入力する", expanded=False):
        st.markdown(
            '<div class="timetable-note">30コマをまとめて設定できます。</div>',
            unsafe_allow_html=True,
        )

        periods = [f"{i}限" for i in range(1, 7)]
        for key in DAY_KEYS:
            cols = st.columns([0.55, 1, 1, 1, 1, 1, 1])
            with cols[0]:
                st.markdown(f"**{DAY_LABELS[key]}**")
            for i, period in enumerate(periods):
                with cols[i + 1]:
                    st.session_state.timetable_input[key][i] = st.text_input(
                        period,
                        value=st.session_state.timetable_input[key][i],
                        key=f"tt_{key}_{i}",
                        placeholder=period,
                        label_visibility="collapsed",
                    )

    # -------------------------------------------------------------------------
    # Advanced options
    # -------------------------------------------------------------------------
    with st.expander("詳細設定", expanded=False):
        c1, c2 = st.columns([1.4, 1])
        with c1:
            selected_template = st.selectbox(
                "Wordテンプレート",
                templates,
                index=0 if templates else None,
                placeholder="テンプレートがありません",
            )
        with c2:
            use_web_search = st.checkbox(
                "Web検索で情報を補足",
                value=False,
                help="必要な場合のみGeminiのWeb検索を使用します。",
            )

        if api_key:
            st.markdown(
                '<span class="status-pill"><span class="status-dot"></span>AI接続済み</span>',
                unsafe_allow_html=True,
            )
        else:
            st.warning("APIキーが設定されていません。環境変数またはStreamlit Secretsを確認してください。")

    # -------------------------------------------------------------------------
    # Generate
    # -------------------------------------------------------------------------
    st.markdown('<div class="section-title">4. お便りを作成</div>', unsafe_allow_html=True)

    st.markdown('<div class="generate-button">', unsafe_allow_html=True)
    generate = st.button(
        "✦  お便りを生成する",
        type="primary",
        disabled=not api_key,
        use_container_width=True,
    )
    st.markdown("</div>", unsafe_allow_html=True)

    if generate:
        if not draft_notes.strip():
            st.error("伝えたい内容を入力してください。")
        else:
            timetable = timetable_from_state()

            with st.status("お便りを作成しています…", expanded=True) as status:
                try:
                    prompt = f"""
あなたはベテラン教員です。以下の条件で、教員が保護者・生徒へ配布するお便りを作成してください。

【条件】
- 読者: {AUDIENCE_OPTIONS[audience]}
- 文書の種類: {DOCUMENT_TYPE_OPTIONS[doc_type]}
- 目安の分量: {LENGTH_OPTIONS[length]}
- 文体: ですます調。温かく前向き。
- 一文をできるだけ短くし、読みやすくする。
- 構成: greeting → body → parent_note → closing
- bodyは具体的な子供たちの様子や言葉を入れ、自然な文章にする。
- 入力された時間割は、mon_1〜fri_6へそのまま反映する。
- 不明な情報を勝手に作りすぎない。

【時間割】
{timetable_text(timetable)}

【教員のメモ】
{draft_notes}
"""

                    data = generate_newsletter_data(api_key, prompt, use_web_search)
                    st.session_state.doc_data = data
                    status.update(
                        label="完成しました。",
                        state="complete",
                        expanded=False,
                    )

                except Exception as error:
                    status.update(
                        label="生成に失敗しました。",
                        state="error",
                        expanded=True,
                    )
                    st.error(f"エラー詳細: {error}")

    # -------------------------------------------------------------------------
    # Preview
    # -------------------------------------------------------------------------
    if st.session_state.doc_data:
        data = st.session_state.doc_data

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown('<div class="section-title">完成したお便り</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="section-caption">内容を確認して、Wordファイルとして保存できます。</div>',
            unsafe_allow_html=True,
        )

        with st.container():
            st.markdown(
                f"""
                <div class="preview-paper">
                    <div class="preview-title">{data.title}</div>
                    <div class="preview-subtitle">
                        {data.subtitle}　·　{data.date_str}
                    </div>
                    <div class="preview-heading">ごあいさつ</div>
                    <div class="preview-body">{data.greeting}</div>
                    <div class="preview-heading">本文</div>
                    <div class="preview-body">{data.body}</div>
                    <div class="preview-heading">保護者の皆様へ</div>
                    <div class="preview-body">{data.parent_note}</div>
                    <div class="preview-heading">結び</div>
                    <div class="preview-body">{data.closing}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        st.markdown("")

        if data.urls:
            with st.expander("参考にした情報"):
                for url in data.urls:
                    st.markdown(f"- {url}")

        if selected_template:
            try:
                docx_bytes = render_docx(selected_template, data)
                st.download_button(
                    label="↓  Wordファイルを保存する",
                    data=docx_bytes,
                    file_name=f"お便り_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as error:
                st.error(f"テンプレートへの埋め込み中にエラーが発生しました: {error}")
        else:
            st.error("templatesフォルダにWordテンプレートがありません。")


if __name__ == "__main__":
    main()
